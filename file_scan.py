#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import re
import xlrd
from docx import Document


def scan_files(directory, prefix=None, postfix=None):
    files_list = []

    for root, sub_dirs, files in os.walk(directory):
        for special_file in files:
            if postfix:
                if special_file.endswith(postfix):
                    files_list.append(os.path.join(root, special_file))
            elif prefix:
                if special_file.startswith(prefix):
                    files_list.append(os.path.join(root, special_file))
            else:
                files_list.append(os.path.join(root, special_file))

    return files_list

def read_text(file_name, keyword):
    with open(file_name) as fn:
        data = fn.readlines()
        pattern = re.compile(keyword.encode('utf-8'))
        for line in data:
            line_word = line.decode('gbk', 'ignore').encode('utf-8')
            if re.search(pattern, line_word):
                return True
    return False

# excel
def read_excel(file_name, keyword):
    book = xlrd.open_workbook(file_name)
    sheet_names = book.sheet_names()
    pattern = re.compile(keyword)
    for sheet_name in sheet_names:
        sheet = book.sheet_by_name(sheet_name)
        nrows = sheet.nrows
        ncols = sheet.ncols
        for rowindex in range(nrows):
            for colindex in range(0, ncols):
                val = sheet.cell(rowindex, colindex).value
                if val:
                    if re.search(pattern, val if val is unicode else u'{}'.format(val)):
                        return True
    return False

# word
def read_word(file_name, keyword):
    if re.search('~\$', file_name):
        return False
    document = Document(file_name)
    pattern = re.compile(keyword)
    for paragraph in document.paragraphs:
        if re.search(pattern, paragraph.text):
            return True
    for table in document.tables:
        for row_idx in range(0, len(table.rows)):
            for cell in table.row_cells(row_idx):
                if re.search(pattern, cell.text):
                    return True
    return False

lst = scan_files(u'D:\\')
for f in lst:
    if f.endswith('.docx'):
        if read_word(f, u'test'):
            print f
    elif f.endswith('.xls') or f.endswith('.xlsx'):
        if read_excel(f, u'test'):
            print f
    else:
        if read_text(f, u'test'):
            print f
